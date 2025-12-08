"""
Serviço de geração de manuscritos em PDF e DOCX
"""
import io
from datetime import datetime
from typing import Optional, List
from django.utils import timezone

# ReportLab para PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.colors import HexColor, grey
from reportlab.pdfgen import canvas

# python-docx para DOCX
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from new_authors.models import AuthorBook, Chapter, PublisherProfile


class WatermarkCanvas(canvas.Canvas):
    """Canvas customizado para adicionar watermark em cada página do PDF"""

    def __init__(self, *args, publisher_name: str = "", **kwargs):
        super().__init__(*args, **kwargs)
        self.publisher_name = publisher_name
        self.page_count = 0

    def showPage(self):
        self.page_count += 1
        self.draw_watermark()
        self.draw_footer()
        super().showPage()

    def draw_watermark(self):
        """Desenha watermark diagonal em cada página"""
        self.saveState()

        # Watermark diagonal no centro
        self.setFont('Helvetica', 40)
        self.setFillColorRGB(0.9, 0.9, 0.9, alpha=0.3)
        self.translate(A4[0] / 2, A4[1] / 2)
        self.rotate(45)

        watermark_text = f"PREVIEW - {self.publisher_name}"
        text_width = self.stringWidth(watermark_text, 'Helvetica', 40)
        self.drawString(-text_width / 2, 0, watermark_text)

        self.restoreState()

    def draw_footer(self):
        """Desenha rodapé com informações da plataforma"""
        self.saveState()

        self.setFont('Helvetica', 8)
        self.setFillColor(grey)

        # Informação à esquerda
        footer_text = f"Baixado por: {self.publisher_name} - {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        self.drawString(2*cm, 1.5*cm, footer_text)

        # Número da página à direita
        page_text = f"Página {self.page_count}"
        self.drawRightString(A4[0] - 2*cm, 1.5*cm, page_text)

        # Plataforma no centro
        platform_text = "CG.BookStore - Plataforma de Talentos Literários"
        text_width = self.stringWidth(platform_text, 'Helvetica', 8)
        self.drawString((A4[0] - text_width) / 2, 1*cm, platform_text)

        self.restoreState()


class ManuscriptGenerator:
    """Gerador de manuscritos em PDF e DOCX"""

    def __init__(self, book: AuthorBook, publisher: PublisherProfile):
        self.book = book
        self.publisher = publisher
        self.author = book.author

    def generate_pdf(
        self,
        chapters: Optional[List[Chapter]] = None,
        full_book: bool = False
    ) -> io.BytesIO:
        """
        Gera PDF do manuscrito com watermark

        Args:
            chapters: Lista de capítulos específicos (None = todos)
            full_book: Se True, inclui todos os capítulos publicados

        Returns:
            BytesIO com o PDF gerado
        """
        buffer = io.BytesIO()

        # Criar documento com canvas customizado
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=3*cm,
            bottomMargin=3*cm,
            title=self.book.title,
            author=self.author.user.get_full_name() or self.author.user.username
        )

        # Estilos
        styles = getSampleStyleSheet()

        # Estilo para título do livro
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        # Estilo para autor
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=14,
            textColor=HexColor('#666666'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Oblique'
        )

        # Estilo para sinopse
        synopsis_style = ParagraphStyle(
            'Synopsis',
            parent=styles['Normal'],
            fontSize=11,
            textColor=HexColor('#333333'),
            alignment=TA_JUSTIFY,
            spaceAfter=30,
            leftIndent=1*cm,
            rightIndent=1*cm
        )

        # Estilo para título de capítulo
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=HexColor('#1a1a1a'),
            spaceAfter=20,
            spaceBefore=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        # Estilo para conteúdo
        content_style = ParagraphStyle(
            'Content',
            parent=styles['BodyText'],
            fontSize=12,
            textColor=HexColor('#000000'),
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=18,
            fontName='Times-Roman'
        )

        # Construir documento
        story = []

        # Página de rosto
        story.append(Spacer(1, 4*cm))
        story.append(Paragraph(self.book.title, title_style))

        if self.book.subtitle:
            story.append(Paragraph(self.book.subtitle, author_style))

        author_name = self.author.user.get_full_name() or self.author.user.username
        story.append(Paragraph(f"por {author_name}", author_style))

        story.append(Spacer(1, 2*cm))

        # Sinopse
        if self.book.synopsis:
            story.append(Paragraph("<b>SINOPSE</b>", synopsis_style))
            story.append(Paragraph(self.book.synopsis, synopsis_style))

        story.append(PageBreak())

        # Informações do manuscrito
        info_style = ParagraphStyle(
            'Info',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#666666'),
            spaceAfter=8
        )

        story.append(Paragraph("<b>INFORMAÇÕES DO MANUSCRITO</b>", chapter_title_style))
        story.append(Spacer(1, 0.5*cm))

        story.append(Paragraph(f"<b>Gênero:</b> {self.book.get_genre_display()}", info_style))
        story.append(Paragraph(f"<b>Idioma:</b> {self.book.get_language_display()}", info_style))
        story.append(Paragraph(f"<b>Status:</b> {self.book.get_status_display()}", info_style))

        if self.book.estimated_pages:
            story.append(Paragraph(f"<b>Páginas estimadas:</b> {self.book.estimated_pages}", info_style))

        story.append(Paragraph(
            f"<b>Visualizações:</b> {self.book.views_count:,}",
            info_style
        ))

        if self.book.rating_average > 0:
            story.append(Paragraph(
                f"<b>Avaliação:</b> {self.book.rating_average:.1f}/5 ({self.book.rating_count} avaliações)",
                info_style
            ))

        story.append(PageBreak())

        # Capítulos
        if chapters is None:
            if full_book:
                chapters = self.book.chapters.filter(is_published=True).order_by('number')
            else:
                chapters = []

        for chapter in chapters:
            # Título do capítulo
            chapter_full_title = f"Capítulo {chapter.number}"
            if chapter.title:
                chapter_full_title += f" - {chapter.title}"

            story.append(Paragraph(chapter_full_title, chapter_title_style))
            story.append(Spacer(1, 0.5*cm))

            # Conteúdo do capítulo
            paragraphs = chapter.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para, content_style))

            # Notas do autor
            if chapter.author_notes:
                story.append(Spacer(1, 1*cm))
                notes_style = ParagraphStyle(
                    'Notes',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=HexColor('#666666'),
                    alignment=TA_JUSTIFY,
                    leftIndent=1*cm,
                    rightIndent=1*cm,
                    borderColor=HexColor('#cccccc'),
                    borderWidth=1,
                    borderPadding=10
                )
                story.append(Paragraph(f"<b>Nota do autor:</b> {chapter.author_notes}", notes_style))

            story.append(PageBreak())

        # Construir PDF com watermark
        doc.build(
            story,
            canvasmaker=lambda *args, **kwargs: WatermarkCanvas(
                *args,
                publisher_name=self.publisher.company_name,
                **kwargs
            )
        )

        buffer.seek(0)
        return buffer

    def generate_docx(
        self,
        chapters: Optional[List[Chapter]] = None,
        full_book: bool = False
    ) -> io.BytesIO:
        """
        Gera DOCX do manuscrito com watermark

        Args:
            chapters: Lista de capítulos específicos (None = todos)
            full_book: Se True, inclui todos os capítulos publicados

        Returns:
            BytesIO com o DOCX gerado
        """
        doc = Document()

        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Adicionar watermark no rodapé
        footer = sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"PREVIEW - {self.publisher.company_name} | Baixado em {datetime.now().strftime('%d/%m/%Y %H:%M')} | CG.BookStore"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Página de rosto
        title = doc.add_heading(self.book.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.book.subtitle:
            subtitle = doc.add_paragraph(self.book.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.italic = True

        author_name = self.author.user.get_full_name() or self.author.user.username
        author_para = doc.add_paragraph(f"por {author_name}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.runs[0]
        author_run.font.size = Pt(12)
        author_run.font.italic = True

        doc.add_paragraph()  # Espaço

        # Sinopse
        if self.book.synopsis:
            doc.add_heading('SINOPSE', level=2)
            synopsis_para = doc.add_paragraph(self.book.synopsis)
            synopsis_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

        # Informações do manuscrito
        doc.add_heading('INFORMAÇÕES DO MANUSCRITO', level=1)

        info_items = [
            f"Gênero: {self.book.get_genre_display()}",
            f"Idioma: {self.book.get_language_display()}",
            f"Status: {self.book.get_status_display()}",
        ]

        if self.book.estimated_pages:
            info_items.append(f"Páginas estimadas: {self.book.estimated_pages}")

        info_items.append(f"Visualizações: {self.book.views_count:,}")

        if self.book.rating_average > 0:
            info_items.append(
                f"Avaliação: {self.book.rating_average:.1f}/5 ({self.book.rating_count} avaliações)"
            )

        for item in info_items:
            para = doc.add_paragraph(item, style='List Bullet')
            para_run = para.runs[0]
            para_run.font.size = Pt(11)

        doc.add_page_break()

        # Capítulos
        if chapters is None:
            if full_book:
                chapters = self.book.chapters.filter(is_published=True).order_by('number')
            else:
                chapters = []

        for chapter in chapters:
            # Título do capítulo
            chapter_full_title = f"Capítulo {chapter.number}"
            if chapter.title:
                chapter_full_title += f" - {chapter.title}"

            chapter_heading = doc.add_heading(chapter_full_title, level=1)
            chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Conteúdo do capítulo
            paragraphs = chapter.content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para_run = para.runs[0]
                    para_run.font.size = Pt(12)
                    para_run.font.name = 'Times New Roman'

            # Notas do autor
            if chapter.author_notes:
                doc.add_paragraph()  # Espaço
                notes_heading = doc.add_paragraph("Nota do autor:")
                notes_run = notes_heading.runs[0]
                notes_run.font.bold = True
                notes_run.font.size = Pt(10)
                notes_run.font.color.rgb = RGBColor(102, 102, 102)

                notes_para = doc.add_paragraph(chapter.author_notes)
                notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                notes_run = notes_para.runs[0]
                notes_run.font.size = Pt(10)
                notes_run.font.color.rgb = RGBColor(102, 102, 102)

            doc.add_page_break()

        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def get_filename(self, file_format: str, chapter: Optional[Chapter] = None) -> str:
        """Gera nome do arquivo para download"""
        from django.utils.text import slugify

        book_slug = slugify(self.book.title)

        if chapter:
            filename = f"{book_slug}_capitulo_{chapter.number}"
        else:
            filename = f"{book_slug}_completo"

        filename += f"_{self.publisher.company_name[:20]}"
        filename = slugify(filename)
        filename += f".{file_format}"

        return filename
